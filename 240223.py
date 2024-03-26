# 파이썬 교육 2일차 ----------------------------------------------------

# 주로 많이 쓰이는 형변환 함수
# int(data)
# str(data)
# list(data)

# if(true) : 참
# print("hi")

#a = "hi python"
#a1 = a[:2]
#a2 = a[3:]
#(a2)

#\\\b = "신현서 입니다."

#b1 = b[4::1]
#b2 = b[0:7:2]

#print(b1)
#print(b2)

#a = [1, 2]

#a.append(3)

#print(a)

#import random
#print(random.randrange(1,46))

#for i in range(6):
#    pass

#lotto - []
#lotto.append()
#lotto.sort()

#not in, sort()


#ctrl + / : 주석



# window 버튼 활용
from genericpath import isfile
import tkinter                                          #pip install tkinter
import time
import random

# E-mail
import os
import smtplib
from email.mime.multipart import MIMEMultipart          # 안 될 경우pip install email
from email.mime.text import MIMEText
from email.mime.base import MIMEBase
from email.mime.image import MIMEImage
from email.mime.application import MIMEApplication
from email import encoders
#- - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - -
# 수료증(워드, 엑셀)
from docx import Document                               #pip install python-docx    #워드 사용
from docx.shared import Pt                              #워드 글짜 크기, 컬러 셋팅
from docx.oxml.ns import qn                             #워드 셋팅
from openpyxl import load_workbook                      #pip install openpyxl       #엑셀 사용
from docx2pdf import convert                            #pip install docx2pdf       #pdf 변환(버전에 따라 안되는 경우 있음)
#- - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - -

num = [0, 0, 0, 0, 0]
My_name = "신현서"

import random

def get_lotto_num():
    lotto = []                            #list로 변수 선언
    while len(lotto) < 6:                 #6개 번호 추출
        number = random.randint(1,45)     #1~45 번호 추출
        if number not in lotto:           #중복 제거, numvers list 내 숫자가 없으면,
            lotto.append(number)          # numbers list에 숫자 삽입
    lotto.sort()                          #6개 값 정렬
    return lotto                          #로또 번호 리턴

def lotto():
    print("이번주 로또 번호는?")
    for i in range(1, 5+1):
        print(get_lotto_num())

def Give_certificates(word, excel, name_tag):
    j = 0

    for i in range(excel.max_row):
        name = excel.cell(row=i+1, column=1).value
        Department = excel.cell(row=i+1, column=2).value

        if i != 0:
            p = word.paragraphs[2]
            if name == name_tag:
                j += 1
                p.text = "성    명: " + name

                #부서
                p = word.paragraphs[3]
                p.text = "부    서: " + Department
                
                #저장
                word.save(f"교육수료증_{name}.docx")

                try:
                    word_name = f"교육수료증_{name}.docx"
                    PDF_name = f"교육수료증_{name}.pdf"
                    convert(word_name, PDF_name)
                    b_pdf = 1
                except:
                    b_pdf = 0

    if j == 0:
        print(f"수료증대상자가 아닙니다.")

    return b_pdf


def certificates():
    doc = Document("교육수료증.docx")
    wb = load_workbook("지원자명단.xlsx")
    try:
        style = doc.styles["Normal"]
        style.font.name = "맑은 고딕"
        style.font.size = Pt(15)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), style.font.name)
    
    except:
        pass

    ws = wb.active
    Give_certificates(doc, ws, My_name)


def naver_email(name, b_file):

    # 비밀번호 입력 파싱
    print("비밀번호를 입력해주세요.")
    id = input("ID : ")
    pwd = input("PW: ")


    # 수료증 확인
    if os.path.isfile(f"./교육수료증_{name}.pdf"):
        target_filename = '교육수료증_%s.pdf' % name

    elif os.path.isfile(f"./교육수료증_{name}.docx"):
        target_filename = '교육수료증_%s.docx' % name

    else:
        print(f"{name}의 교육 수료증이 존재 하지 않습니다.")
        return()

    try:
        # smtp 이용(Simple Mail Transfer Protocol) : 인터넷을 이용하여 이메일을 보낼 때 사용하는 프로토콜
        mail_server = smtplib.SMTP('smtp.naver.com', 587)
        mail_server.ehlo()
        mail_server.starttls()
        mail_server.ehlo()

        # 로그인
        mail_server.login(id, pwd)
    except:
        print(f"로그인 오류!")
        return()

    # 제목, 본문 작성
    msg = MIMEMultipart()  
    # 보내는 사람
    msg['From'] = 'hyunseo1018@naver.com'
    # 받은 사람
    msg['To'] = "hyunseo1018@nicepay.co.kr"
    # 메일 제목   
    msg['Subject'] = f"{name} 님 ZICO 교육 수료증 입니다."
    # 메일 내용
    msg.attach(MIMEText('zico교육 모두수고하셨습니다.', 'plain'))
    # 파일첨부 (파일 미첨부시 생략가능)
    if b_file == True:
        attachment = open('./%s' % target_filename, 'rb')
        part = MIMEBase('application', 'octet-stream')
        part.set_payload(attachment.read())
        encoders.encode_base64(part)
        file = os.path.basename('./%s'%target_filename)
        part.add_header('Content-Disposition', "attachment", filename=file)
        msg.attach(part)

    # 메일 전송
    re = mail_server.sendmail(msg['From'], msg['To'].split(','), msg.as_string())
    if re == {}:
        print(f"메일 정상 전송 되었습니다.")
    else:
        print(f"메일 전송 오류!")
    mail_server.quit()

def event1():
    num[0] += 1
    bt1["text"] = f"로또 {num[0]}번"
    lotto()

def event2():
    num[1] += 1
    bt2["text"] = f"교육수료증 {num[1]}번"
    certificates()
def event3():
    num[2] += 1
    bt3["text"] = f"이메일 전송 {num[2]}번"
    naver_email(My_name, True)

def event4():
    num[3] += 1
    bt4["text"] = f"GUI로또 {num[3]}번"
    lotto()

def event5():
    num[4] += 1
    bt5["text"] = f"5버튼누름 {num[4]}번"
    data = enter.get()
    label["text"] = data

def event6():
    global num
    bt1["text"] = "  로또  "
    bt2["text"] = "  교육수료증  "
    bt3["text"] = " 이메일(ID/PW 입력) "
    bt4["text"] = "  GUI로또  "
    bt5["text"] = "  버튼오  "
    enter.delete(0, "end")
    num = [0, 0, 0, 0, 0]

tk = tkinter.Tk()
tk.title("파이썬 GUI 프로그램")
tk.geometry("400x600")                          # 창의 크기 ( W * H )

enter = tkinter.Entry(tk)                       # 값 입력
enter.pack(side=tkinter.TOP, padx = 18)         # enter 창 위치(상대좌표)
# enter.place(x=1, y=1)                         # enter 창 위치(절대좌표)

label = tkinter.Label(tk, text = " [ 문구표시 ] ")
label.pack(side=tkinter.BOTTOM)

# 버튼생성
bt1 = tkinter.Button(tk, text = "  로또  ", command=event1, width=18)
bt2 = tkinter.Button(tk, text = "  교육수료증  ", command=event2, width=18)
bt3 = tkinter.Button(tk, text = " 이메일(ID/PW 입력) ", command=event3, width=18)
bt4 = tkinter.Button(tk, text = " GUI 로또 ", command=event4, width=18)
bt5 = tkinter.Button(tk, text = "  버튼오  ", command=event5, width=18)
bt6 = tkinter.Button(tk, text = "  초기화  ", command=event6, width=18)

# 버튼위치
bt1.pack(side=tkinter.TOP, padx = 10, pady = 4)
bt2.pack(side=tkinter.TOP, padx = 10, pady = 4)
bt3.pack(side=tkinter.TOP, padx = 10, pady = 4)
bt4.pack(side=tkinter.TOP, padx = 10, pady = 4)
bt5.pack(side=tkinter.TOP, padx = 10, pady = 4)
bt6.pack(side=tkinter.TOP, padx = 10, pady = 4)

tk.mainloop()